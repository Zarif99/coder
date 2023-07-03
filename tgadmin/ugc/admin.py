import base64
import hashlib
import os
import re
from functools import reduce
from io import BytesIO
from pathlib import Path
from typing import Union, Optional
from urllib.parse import urlparse, parse_qs

import docx.text.paragraph
import requests
from PIL import Image as PIL_Image
from django.conf import settings
from docx import oxml, opc, Document  # Import the Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK_TYPE
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement, parse_xml, nsmap
from docx.oxml.ns import qn, nsdecls
from docx.shared import RGBColor, Pt, Cm, Inches  # Import Pt, used to font etc

from docsie.api.models import Documentation, File, Book, Snippet, Article
from docsie.common.utils import capture_exception
from docsie.files.uploader.document import DocumentUploadProvider
from docsie.users.models import User

try:
    from cairosvg import svg2png
except OSError:
    pass
    #  Don't installed https://doc.courtbouillon.org/weasyprint/stable/first_steps.html
    #  DOCX Import an .svg image will no works.

from .source.xml_templates import add_float_picture

raven = settings.RAVEN

HEADERS = {
    'header-two': 2,
    'header-three': 3
}

PATH_TO_SOURCE = Path(Path(__file__).parent.resolve(), 'source')

LISTS = {
    'ordered-list-item': 'List Number',
    'unordered-list-item': 'List Bullet'
}
NAMESPACES = {
    'o': 'urn:schemas-microsoft-com:office:office',
    've': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w10': 'urn:schemas-microsoft-com:office:word',
    'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'mv': 'urn:schemas-microsoft-com:mac:vml',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'mo': 'http://schemas.microsoft.com/office/mac/office/2008/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'v': 'urn:schemas-microsoft-com:vml',
    'wp': ('http://schemas.openxmlformats.org/drawingml/2006/wordprocessing'
           'Drawing'),
    'cp': ('http://schemas.openxmlformats.org/package/2006/metadata/core-pr'
           'operties'),
    'dc': 'http://purl.org/dc/elements/1.1/',
    'ep': ('http://schemas.openxmlformats.org/officeDocument/2006/extended-'
           'properties'),
    'xsi': 'http://www.w3.org/2001/XMLSchema-instance',
    'ct': 'http://schemas.openxmlformats.org/package/2006/content-types',
    'r': ('http://schemas.openxmlformats.org/officeDocument/2006/relationships'),
    'pr': 'http://schemas.openxmlformats.org/package/2006/relationships',
    'dcmitype': 'http://purl.org/dc/dcmitype/',
    'dcterms': 'http://purl.org/dc/terms/',
    'wp15': 'http://schemas.microsoft.com/office/word/2012/wordprocessingDrawing',
    'a14': 'http://schemas.microsoft.com/office/drawing/2010/main',
}


nsmap.update(NAMESPACES)

DOCUMENT_NS = '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:mv="urn:schemas-microsoft-com:mac:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" mc:Ignorable="w14 wp14">\n<w:body>'
SET_XML_INTO_NAMESPACES = lambda xml: '{}{}</w:body></w:document>'.format(DOCUMENT_NS,
                                                                          xml)  # for creating XML format which lxml can parse


def set_italic(run):
    try:
        run.italic = True
        return run
    except Exception:
        raven.captureException()


def set_bold(run, ):
    try:
        run.bold = True
        return run
    except Exception:
        raven.captureException()


def set_underline(run):
    try:
        run.underline = True
        return run
    except Exception:
        raven.captureException()


def set_link(run, self, href):
    try:
        self._add_link_by_run(run=run, href=href)
    except Exception:
        raven.captureException()


def set_styles(name, run, self=None, text=''):
    try:
        if 'link' in name:
            run = self._add_link_by_run(run, name['link'], text=text)
        elif 'img' in name:
            run.text = run.text.replace('🖼', '', 1)  # removing symbol 🖼 (image symbol)
            image_content = requests.get(name['img']['data']['src']).content
            size = int(name['img']['data']['size'])

            # it needs to convert pixels to Inches:
            size = size / 96
            # You can see about converting pixels to inches here:
            # https://www.codegrepper.com/code-examples/whatever/pixel+to+inches

            size = Inches(size)
            run.add_picture(BytesIO(image_content), width=size, height=size)
        elif name == 'header-step':
            run.bold = True
            run.font.size = Pt(10.5)
        elif name == 'BOLD':
            set_bold(run)
        elif name == 'ITALIC':
            set_italic(run)
        elif name == 'UNDERLINE':
            set_underline(run)
        elif name == 'CODE':
            run.font.color.rgb = RGBColor(68, 114, 196)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        elif name == 'KBD':
            tag = run._r  # get XML of run
            shd = OxmlElement('w:shd')  # create a shadow to will will it to gray
            shd.set(qn('w:fill'), 'e7e6e6')  # fill to gray background block code
            run.font.size = Pt(11)
            run.font.name = 'Courier New'
            tag.rPr.append(shd)  # append to XML of run shadow
        elif name == 'DFN':
            tag = run._r  # get XML of run
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '#b3c6e7')
            tag.rPr.append(shd)  # append to XML of run shadow
        return run
    except Exception as e:
        raven.captureException()


def remove_crap(craplist, st):
    try:
        for c in craplist:
            st = st.replace(c, '')
        return st
    except Exception:
        raven.captureException()


class DocxExporter:
    def __init__(self, shelf_data: dict, based_template: Optional[Union[Document, BytesIO, bytes]] = None):
        try:
            self.shelf = shelf_data

            # General
            self.document: Document = Document()
            self._paragraph = None
            self.__depth_level = 1
            self.section = self.document.sections[0]
            self.images_max_width = Cm(
                14.8)  # base width is 15, but for more look better around the edges I set that to 14.8
            # that is default font settings

            self.__font_size = Pt(12)
            self.__font_name = 'Inter'
            self.__font_color_rgb = RGBColor(64, 64, 64)
            self.__header_font_color_rgb = RGBColor(52, 171, 118)

            # Table
            self.__table = None
            self.__table_col_index = 0
            self.__cols_num = 0

            self.cells_d = 0
            self.rows_d = 0
            # merge
            self.__based_template = None

            if based_template:
                if type(based_template) == Document:
                    self.__based_template = based_template
                elif type(based_template) == BytesIO:
                    self.__based_template = Document(based_template)
                elif type(based_template) == bytes:
                    self.__based_template = Document(BytesIO(based_template))
                else:
                    pass
        except Exception:
            raven.captureException()

    def export(self):
        try:
            # self.add_ToC_element()
            self._add_title_page(self.shelf['shelf_name'])
            for i, book in enumerate(self.shelf['books'], start=0):
                try:
                    book_name = book['name']  # set name of book to the variable
                    articles = book['articles']  # set articles of book to the variable
                    self._add_book_title(book_name)

                    self.__depth_level = 1
                    self.step_header_elems = {'num': 0}  # this is for header step element
                    self.step_list_items = {'num': 0}  # this is for list items
                    self.cells = 0  # set empty cells
                    self.rows = 0  # set empty rows
                    self.__dict = None  # set empty dict

                    for article in articles:
                        # self.__depth_level = 2
                        self.entity_map = article['entity_map']  # set entitymap to global variable in the class

                        parsed_blocks = []
                        for block in article['blocks']:
                            if block['type'] == 'snippet':
                                snippet = Snippet.objects.select_related('article').get(id=block['data']['src'])
                                snippet_blocks = snippet.blocks
                                source_article = Article.objects.get(id=snippet.article_id)
                                parsed_blocks.extend(
                                    [block for block in source_article.doc['blocks'] if block['key'] in snippet_blocks])
                            else:
                                parsed_blocks.append(block)

                        self.article_blocks = parsed_blocks  # set article blocks to global variable in the class
                        self.keys = [key['key'] for key in self.article_blocks]  # get keys from blocks
                        self.doc_version = article['doc_version']
                        self._set_article(article)
                        try:
                            if self.doc_version <= 2:
                                index_cell = 0
                                while index_cell < len(self.article_blocks):  #
                                    if self.article_blocks[index_cell][
                                        'type'] == 'cell':  # if item is cell, need to configure depths
                                        depths = []
                                        try:
                                            for index_block in range(0, 3):  # get depths of cells items
                                                depths.append(
                                                    self.article_blocks[
                                                        self.keys.index(
                                                            self.article_blocks[index_cell]['key']) + index_block][
                                                        'depth'])
                                        except IndexError:  # if not all is cells, or table < 3 columns
                                            depths = [1, 1,
                                                      1]  # create custom depths and set value [1,1,1], to next check didn't not pass, because current table < 3 columns
                                        if self.article_blocks[index_cell]['depth'] == 1 and depths[
                                            2] == 2:  # configure strage system of depths to normal (if table have 3 columns)
                                            self.article_blocks[index_cell]['depth'] = 3
                                            self.article_blocks[index_cell + 1]['depth'] = 3
                                            self.article_blocks[index_cell + 2]['depth'] = 3
                                            index_cell += 1
                                    index_cell += 1
                        except:  # if function to set cells will return error that except resolve this
                            pass

                        self.article_blocks.append(
                            {'type': 'unstyled', 'offset': 0, 'length': 0, 'text': '', 'depth': 0, 'entityRanges': [],
                             'inlineStyleRanges': []})
                        for block in self.article_blocks:  # for block in the article blocks, on here that 'for' get styles of blocks, and append elements depending on type
                            try:
                                if block['type'] == 'unstyled' and [enty for enty in block['entityRanges'] if
                                                                    self.entity_map.get(str(enty.get('key')), {}).get(
                                                                        'data', {}).get('style') == 'block']:
                                    self.add_block_link(block)
                                elif block[
                                    'type'] == 'unstyled':  # if this is unstyled text, need only create a new paragraph, and add text
                                    self.__reset_non_general()
                                    self.__add_paragraph(self.document)
                                    self.__depth_level += block.get('depth', 0)
                                    self.add_text(block=block)
                                    self.__depth_level -= block.get('depth', 0)


                                elif block['type'] == 'figure':  # if this is picture, need to add fugure (picture)
                                    self._add_figure(block)

                                elif block['type'] == 'mdtable':
                                    self._add_mdtable(block=block)

                                elif block['type'] in HEADERS.keys():  # if this is anything header, need to add heading
                                    self._add_headers(block=block)

                                elif block[
                                    'type'] in LISTS.keys():  # if this is anything list item (bullet or unordered), need to append list item
                                    self._add_list_item(block)


                                elif block[
                                    'type'] == 'dictionary':  # if this is dictionary, need to append table (dictionary)
                                    self._add_dict(block)

                                elif block['type'] == 'cell':  # if this is cell need to append cell or create new table
                                    self._add_table_cell(block)

                                elif block[
                                    'type'] == 'blockquote':  # if this is blockquote, need to append a blockquote item
                                    # We use another condition for blockquote in order to style it in future
                                    self._add_blockquote(block, block['data'].get('style'))

                                elif block[
                                    'type'] == 'header-step':  # if this is header-step, need to append a header-step
                                    self._add_header_step(block)

                                elif block[
                                    'type'] == 'code-block':  # if this is code-block, need to append a code-block
                                    self._add_block_code(block)

                                elif block['type'] == 'video':  # if this is video, need to append a video item
                                    self._add_video(block)
                                elif block['type'] == 'gist-block':
                                    self._add_gist(block)

                                else:  # if this is another style, add text (if it is have text), and reset values
                                    self.reset(block)
                                self.__depth_level = 1
                            except Exception as e:
                                raven.captureException()
                except Exception as e:
                    raven.captureException()
                self._paragraph.add_run().add_break(WD_BREAK_TYPE.PAGE)  # end page
                self.__merge__()  # if has based document, that will merge styles
        except Exception as e:
            raven.captureException()

    def _set_article(self, article):
        try:
            self.__add_paragraph(self.document)

            if article.get('meta') and article['meta'].get(
                'icon'):  # if artlice have icon, need to create a picture in the .docx file
                self._set_picture(url=article['meta']['icon'], alignment=1)  # and set article name
            self.__add_paragraph(self.document)
            self._paragraph.add_run('\n')

            run_name = self._paragraph.add_run(article.get('name'))
            run_name.font.size = Pt(18)
            run_name.font.name = self.__font_name
            run_name.font.color.rgb = self.__font_color_rgb
            set_bold(run_name)
            self._paragraph.alignment = 1
            self.__add_paragraph(self.document)
            run_name = self._paragraph.add_run(article.get('description'))
            run_name.font.size = Pt(12)
            run_name.font.name = self.__font_name
            run_name.font.color.rgb = self.__font_color_rgb
            set_bold(run_name)
            self._paragraph.alignment = 1

            self.__add_paragraph(self.document)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '404040')
            color.set(qn('w:themeColor'), 'text1')
            color.set(qn('w:themeTint'), 'BF')

            size = OxmlElement('w:sz')
            size.set(qn('w:val'), '26')

            size_cs = OxmlElement('w:szCs')
            size_cs.set(qn('w:val'), '26')

            fonts = OxmlElement('w:rFonts')
            fonts.set(qn('w:ascii'), 'Inter')
            fonts.set(qn('w:hAnsi'), 'Inter')
            rPr = OxmlElement('w:rPr')

            rPr.append(color)
            rPr.append(size)
            rPr.append(size_cs)

            self._paragraph._p.pPr.append(rPr)
        except Exception:
            raven.captureException()

    def __merge__(self):
        try:
            if self.__based_template:
                self.document = DocxMergeStyles.merge(input_doc=self.document, based_doc=self.__based_template)
            else:
                pass
        except Exception:
            raven.captureException()

    def parse_xml_string(self, content: str):
        try:
            xml = SET_XML_INTO_NAMESPACES(content)
            parsed_xml = parse_xml(xml)
            finally_element = parsed_xml[0][0]  # getting [0][0] cuz needs to remove Document and Body tags
            return finally_element
        except Exception:
            raven.captureException()

    def reset(self, block):

        try:
            if 'text' in block:
                self.add_text(block)
        except Exception:
            raven.captureException()

    def __add_paragraph(self, position, style: str = 'Normal'):
        try:
            self._paragraph = position.add_paragraph(style=style)
            self._paragraph.style.font.name = 'Inter'
            self._paragraph.style.font.color.rgb = self.__font_color_rgb
            if style == 'Caption':  # if style caption, neet to set another style (style for headers)
                self._paragraph.style.font.size = Pt(10)  # set font size of default header
                self._paragraph.style.font.name = 'Inter'
                self._paragraph.style.font.color.rgb = self.__font_color_rgb
                self._paragraph.alignment = 1
        except Exception:
            raven.captureException()

    def _create_oxml_element(self, content: str):
        try:
            finally_xml = SET_XML_INTO_NAMESPACES(content)
            element = oxml.etree.fromstring(finally_xml)  # this element can be used in docx lib
            return element
        except Exception:
            raven.captureException()

    def add_block_link(self, block):
        try:
            self.__add_paragraph(self.document)
            href = ''
            text = ''
            for entity in block['entityRanges']:
                if self.entity_map.get(str(entity.get('key')), {}).get('data', {}).get('style') == 'block':
                    data = self.entity_map.get(str(entity.get('key')), {}).get('data', {})
                    href = data.get('href')
                    text = block['text']

            run_text = self._paragraph.add_run(text + '\n')
            run_text.font.color.rgb = RGBColor(76, 174, 227)

            run_href = self._paragraph.add_run(href)
            set_italic(run_href)
            set_italic(run_text)
            run_href.font.size = Pt(12)
            run_text.font.size = Pt(12)
            run_href.font.name = self.__font_name
            run_text.font.name = self.__font_name
        except Exception:
            raven.captureException()

    def _add_gist(self, block):
        try:
            self.__add_paragraph(self.document)
            block_code_table = self.document.add_table(rows=1, cols=1)
            self.set_cell_border(block_code_table.rows[0].cells[0],
                                 start={'sz': 12, 'color': '#365FDD', 'val': 'single'})
            block_code_row = block_code_table.rows[0]
            block_code_row.cells[0].text = block.get('data', {}).get('src', '')
            block_code_row.cells[0].paragraphs[0].style.name = 'HTML Preformatted'
            block_code_row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            block_code_row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        except Exception:
            raven.captureException()

    def _read_markdown_table(self, table: str):
        result = []
        for n, line in enumerate(table[1:-1].split('\n')):
            data = {}
            if n == 0:
                header = [t.strip() for t in line.split('|')[1:-1]]
            if n > 1:
                values = [t.strip() for t in line.split('|')[1:-1]]
                for col, value in zip(header, values):
                    data[col] = value
                result.append(data)
        return result

    def _add_mdtable(self, block):
        json_table = self._read_markdown_table(block['text'])
        if len(json_table) == 0: return
        rows = list(json_table[0].keys())
        columns = []
        for row in json_table:
            row_values = []
            for column in row:
                row_values.append(row[column])
            columns.append(row_values)

        table = self.document.add_table(rows=0, cols=len(columns[0]))
        table.add_row()
        table.style = 'Table Grid'
        for cell_id, row in enumerate(rows):
            table.rows[- 1].cells[cell_id].text = row

        for columns_ in columns:
            table.add_row()
            row = table.rows[- 1]
            for cell_id, column in enumerate(columns_):
                row.cells[cell_id].text = column

    def upload_image(self, location: str, image_data, content_type: str = 'image/*'):
        try:
            bucket = settings.DOCSIE_UPLOAD_BUCKET_NAME
            resource = settings.S3_RESOURCE
            obj = resource.Object(bucket, location)
            obj.put(Body=image_data, ContentType=content_type, ACL='public-read')
            url = 'https://{bucket}.s3.amazonaws.com/{location}'.format(location=location, bucket=bucket)
            return url
        except Exception:
            raven.captureException()

    def _add_shadow_header(self, text: str):
        """
        This function uses for adding currently name of reading paragraph in the Docsie file.
        """
        self.__add_paragraph(self.document)
        run = self._paragraph.add_run(text)
        self._paragraph.alignment = 2
        run.font.size = Pt(10)
        run.font.name = self.__font_name
        run.font.color.rgb = RGBColor(127, 127, 127)
        run.bold = False

    def is_here_break_page(self, run: docx.text.paragraph.Run = None):
        """
        That function can return end of the page if current Docx document ever has been opened in Word
        """
        if 'lastRenderedPageBreak' in run._element.xml:
            return True
        elif 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
            return True
        return False

    def _add_headers(self, block: dict):
        type_header = block.get('type')

        if type_header == 'header-two':
            self.__add_paragraph(self.document, style='heading 1')
            run_heading = self._paragraph.add_run(block.get('text'))
            run_heading.font.size = Pt(16)
            run_heading.font.color.rgb = RGBColor(52, 171, 118)
            run_heading.font.name = 'Inter'
            set_bold(run_heading)
        elif type_header == 'header-three':
            self.__add_paragraph(self.document)
            run = self._paragraph.add_run(block.get('text', ''))
            run.font.size = Pt(12)
            run.font.name = self.__font_name
            run.font.color.rgb = self.__font_color_rgb
            set_bold(run)

    def _add_heading(self, text: str, position=None, block=None, heading_num: int = 2,
                     alignment: int = 0, font_size: float = None, style: str = None):

        try:
            if not position:
                position = self.document
            if not style:
                style = f'heading {heading_num}'

            self.__add_paragraph(position, style=style)  # add paragraph, and set style of heading
            if block:
                self.add_text(block)  # add text of block
            else:
                run = self._paragraph.add_run(
                    text)
                run.font.name = 'Inter'
            if style:
                self._paragraph.style.name = style
            self._paragraph.style.font.name = 'Inter'
            if not font_size:
                if heading_num == 1:
                    self._paragraph.style.font.size = Pt(16)
                elif heading_num == 2:
                    self._paragraph.style.font.size = Pt(14)
                elif heading_num == 3:
                    self._paragraph.style.font.size = Pt(12)
                else:
                    self._paragraph.style.font.size = Pt(11)
            else:
                self._paragraph.style.font.size = font_size
            self._paragraph.alignment = alignment
        except Exception:
            raven.captureException()

    def _add_book_title(self, book_name: str):
        try:
            run = self._paragraph.add_run(book_name)
            run.font.name = 'Inter'
            run.font.size = Pt(24)
            run.font.color.rgb = self.__font_color_rgb
            set_bold(run)
        except Exception:
            raven.captureException()

    def _add_title_page(self, shelf_name):
        '''
        Function for creating the top of page.
        It's adding shelf or book name at the top and adding background image
        '''
        try:
            self.__add_paragraph(self.document)
            run = self._paragraph.add_run(text=shelf_name)
            run.font.name = 'Inter'
            run.font.size = Pt(28)
            run.font.color.rgb = self.__header_font_color_rgb
            run.bold = True

            self.__add_paragraph(self.document)

            add_float_picture(self._paragraph,
                              str(Path(PATH_TO_SOURCE, 'background.jpg')))

        except Exception:
            raven.captureException()

    def _add_block_code(self, block):  # function to add block code
        try:
            self.__add_paragraph(self.document)  # create a new paragraph
            information_table = self.document.add_table(rows=1, cols=2)
            information_table.rows[0].cells[0].text = block.get('data', {}).get('label', '')
            information_table.rows[0].cells[1].text = block.get('data', {}).get('type', '')
            information_table.rows[0].cells[0].paragraphs[0].runs[0].font.size = self.__font_size
            information_table.rows[0].cells[1].paragraphs[0].runs[0].font.size = self.__font_size
            information_table.rows[0].cells[0].paragraphs[0].runs[0].font.name = self.__font_name
            information_table.rows[0].cells[1].paragraphs[0].runs[0].font.name = self.__font_name
            information_table.columns[0].height = Inches(0.6)

            block_code_table = self.document.add_table(rows=1, cols=1)
            self.set_cell_border(block_code_table.rows[0].cells[0],
                                 start={'sz': 12, 'color': '#365FDD', 'val': 'single'})
            block_code_row = block_code_table.rows[0]
            block_code_row.cells[0].text = block.get('text', '')
            block_code_row.cells[0].paragraphs[0].style.name = 'HTML Preformatted'
            block_code_row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            block_code_row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'

        except Exception as e:
            raven.captureException()

    def __get_video_id_youtube(self, url):
        try:
            '''
            Examples:
            - http://youtu.be/SA2iWivDJiE
            - http://www.youtube.com/watch?v=_oPAwA_Udwc&feature=feedu
            - http://www.youtube.com/embed/SA2iWivDJiE
            - http://www.youtube.com/v/SA2iWivDJiE?version=3&amp;hl=en_US
            '''
            query = urlparse(url)
            if query.hostname == 'youtu.be':
                return query.path[1:]
            if query.hostname in ('www.youtube.com', 'youtube.com'):
                if query.path == '/watch':
                    p = parse_qs(query.query)
                    return p['v'][0]
                if query.path[:7] == '/embed/':
                    return query.path.split('/')[2]
                if query.path[:3] == '/v/':
                    return query.path.split('/')[2]

            return
        except Exception:
            raven.captureException()

    def __get_video_id_vimeo(self, url):
        try:
            cmp = r'https?:\/\/(?:www\.|player\.)?vimeo.com\/(?:channels\/(?:\w+\/)?|groups\/([^\/]*)\/videos\/|album\/(\d+)\/video\/|video\/|)(\d+)(?:$|\/|\?)'
            matches = re.search(cmp, url)
            for res in matches.groups():
                if res:
                    return res
            return
        except Exception:
            raven.captureException()

    def _add_video(self, block):  # function to add video
        '''
        That function create image of video, and link to video in the bottom of image.
        '''
        try:

            self.__add_paragraph(self.document)
            self._paragraph.alignment = 1  # add aligment in the center
            images = {}
            try:
                url = block['data']['src']  # try to get src of video
            except KeyError:  # if not have url
                return
            if 'youtu' in url:  # YouTube video
                try:
                    video_id = self.__get_video_id_youtube(url)

                    images = {
                        'small_image': 'https://img.youtube.com/vi/{}/default.jpg'.format(video_id),
                        # small image of video
                        'normal_image': 'https://img.youtube.com/vi/{}/hqdefault.jpg'.format(video_id),
                        # normal image of video
                        'big_image': 'https://img.youtube.com/vi/{}/sddefault.jpg'.format(video_id)
                        # big image of video

                    }
                except Exception:
                    pass  # incorrect url, or incorrect video, and more problems in YouTube
            elif 'vimeo' in url:  # Vimeo
                try:
                    id_video = self.__get_video_id_vimeo(url)
                    data_from_video = requests.get('http://vimeo.com/api/v2/video/{}.json'.format(id_video)).json()
                    images = {
                        'small_image': data_from_video[0]['thumbnail_small'],  # small image of video
                        'normal_image': data_from_video[0]['thumbnail_medium'],  # normal image of video
                        'big_image': data_from_video[0]['thumbnail_large']  # big image of video
                    }
                except Exception:
                    pass  # incorrect url, or incorrect video, and more problems in Vimeo

            if images.get('big_image'):
                self._set_picture(depth=0, label='', url=images['big_image'], alignment=1, border=True)
            # from lxml import etree
            # from xml.etree import ElementTree as ET
            # pic_cNvPr = None
            # pic_pic = None
            # graphic_data = None
            # wp_inline = None
            # for tag in self._paragraph._p.iter():
            #     print(tag.tag)
            #     if tag.tag == '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr':
            #         pic_cNvPr = tag
            #     elif tag.tag == '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData':
            #         graphic_data = tag
            #     elif tag.tag == '{http://schemas.openxmlformats.org/drawingml/2006/picture}pic':
            #         pic_pic = tag
            #     elif tag.tag == '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline':
            #         wp_inline = tag
            #
            # pic_id = pic_cNvPr.attrib['id']
            # pic_name = pic_cNvPr.attrib['name']
            # pic_pic.getparent().remove(pic_pic)
            #
            # new_pic_pic = OxmlElement('pic:pic')
            #
            # pic_nvPicPr = OxmlElement('pic:nvPicPr')
            # pic_cNvPr = OxmlElement('pic:cNvPr')
            # a_hlinkClick = OxmlElement('a:hlinkClick')
            # pic_blipFill = OxmlElement('pic:blipFill')
            # a_blip = OxmlElement('a:blip')
            # a_extLst = OxmlElement('a:extLst')
            # a_ext1 = OxmlElement('a:ext')
            # a14_useLocalDpi = OxmlElement('a14:useLocalDpi')
            # a_ext2 = OxmlElement('a:ext')
            # wp15_webVideoPr = OxmlElement('wp15:webVideoPr')
            # a_stretch = OxmlElement('a:stretch')
            # a_fillRect = OxmlElement('a:fillRect')
            # pic_spPr = OxmlElement('pic:spPr')
            # a_xfrm = OxmlElement('a:xfrm')
            # a_off = OxmlElement('a:off')
            # a_ext_xfrm = OxmlElement('a:ext')
            # a_prstGeom = OxmlElement('a:prstGeom')
            # a_avLst = OxmlElement('a:avLst')
            #
            # pic_nvPicPr.append(pic_cNvPr)
            # pic_cNvPr.append(a_hlinkClick)
            # pic_blipFill.append(a_blip)
            #
            # a_blip.append(a_extLst)
            # a_extLst.append(a_ext1)
            # a_ext1.append(a14_useLocalDpi)
            # a_extLst.append(a_ext2)
            # a_ext2.append(wp15_webVideoPr)
            #
            # pic_blipFill.append(a_stretch)
            # a_stretch.append(a_fillRect)
            #
            # pic_spPr.append(a_xfrm)
            # a_xfrm.append(a_off)
            # a_xfrm.append(a_ext_xfrm)
            #
            # pic_spPr.append(a_prstGeom)
            # a_prstGeom.append(a_avLst)
            #
            # new_pic_pic.append(pic_cNvPr)
            # new_pic_pic.append(pic_blipFill)
            # new_pic_pic.append(pic_spPr)
            #
            # graphic_data.append(new_pic_pic)
            #
            # pic_cNvPr.set('id',pic_id)
            # pic_cNvPr.set('name',pic_name)
            # pic_cNvPr.set('descr','video')
            #
            # a_ext1.set('uri','{28A0092B-C50C-407E-A947-70E740481C1C}')
            # a14_useLocalDpi.set('val','0')
            # a_ext2.set('uri','{C809E66F-F1BF-436E-b5F7-EEA9579F0CBA}')
            # wp15_webVideoPr.set('embeddedHtml',f'<iframe width="200" height="113" src="{url}?feature=oembed" frameborder="0" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture" allowfullscreen="" sandbox="allow-scripts allow-same-origin allow-popups"></iframe>')
            # wp15_webVideoPr.set('h','113')
            # wp15_webVideoPr.set('w','200')
            # a_off.set('x','0')
            # a_off.set('y','0')
            # a_ext_xfrm.set('cx','4572000')
            # a_ext_xfrm.set('cy','3429000')
            #
            # a_prstGeom.set('prst','rect')
            #
            # # Wp inline
            # wp_docPr = OxmlElement('wp:docPr')
            # a_hlinkClick_inline = OxmlElement('a:hlinkClick ')
            # wp_inline.set('distT','0')
            # wp_inline.set('distB','0')
            # wp_inline.set('distL','0')
            # wp_inline.set('distR','0')
            # wp_docPr.set('id',pic_id)
            # wp_docPr.set('name',pic_name)
            # wp_docPr.set('descr','video')
            # wp_inline.append(wp_docPr)
            # wp_docPr.append(a_hlinkClick_inline)
            #

            label = block['data'].get('label', '')
            self.__add_paragraph(self.document, style='Caption')
            run = self._paragraph.add_run(text=label)
            set_bold(run)
            self._paragraph.add_run(text='\n')
            run = self._paragraph.add_run(text=block['data']['src'])
            set_bold(run)
        except Exception:
            raven.captureException()

    def set_cell_border(self, cell, **kwargs):
        '''
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={'sz': 12, 'val': 'single', 'color': '#FF0000', 'space': '0'},
            bottom={'sz': 12, 'color': '#00FF00', 'val': 'single'},
            start={'sz': 24, 'val': 'dashed', 'shadow': 'true'},
            end={'sz': 12, 'val': 'dashed'},
        )
        '''
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ['sz', 'val', 'color', 'space', 'shadow']:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def add_text(self, block, paragraph=None):
        try:
            if not paragraph:
                paragraph = self._paragraph
            entity_ranges = block.get('entityRanges', [])  # save entity ranges (links) to variable, for
            inline_style_ranges = block.get('inlineStyleRanges', [])  # save entity ranges (links) to variable
            if entity_ranges:
                # set entity's in inline_style_ranges
                for entity_r in range(len(entity_ranges)):  # 'for' to append links in readable format.
                    # create links format to common styles. Because in .docsie format links and styles (bold,italic,etc) is two different formats, that function will set
                    # that links and styles in one format, for example: {'style':{'link'}},{'style':{'bold'}},etc
                    entity_range_id = str(entity_ranges[entity_r].get('key'))
                    if entity_range_id not in self.entity_map or 'type' not in self.entity_map[entity_range_id]:
                        continue
                    entity_type = self.entity_map[entity_range_id]['type']

                    if entity_type == 'LINK':
                        inline_style_ranges.append(
                            {'style': {'link': self.entity_map[str(entity_ranges[entity_r]['key'])]['data']['href']},
                             'offset': entity_ranges[entity_r]['offset'], 'length': entity_ranges[entity_r]['length']})
                    elif entity_type == 'IMG':
                        inline_style_ranges.append(
                            {'style': {'img': self.entity_map[str(entity_ranges[entity_r]['key'])]},
                             'offset': entity_ranges[entity_r]['offset'], 'length': entity_ranges[entity_r]['length']})
            if block[
                'type'] == 'header-step':  # append style header-step. Because header step is different type of cell, dictionary, video, etc
                # because header-step can have styles (bold, italic, etc)
                # styles bold, italic and underline have only in unstyled and header-steps
                inline_style_ranges.append({'style': 'header-step', 'offset': 0, 'length': len(block['text'])})

            # bubble alghoritm to set a words in normal position
            # because after previous 'for', position elements was set incorrect
            for i in range(len(inline_style_ranges) - 1):
                for j in range(len(inline_style_ranges) - i - 1):
                    if inline_style_ranges[j]['offset'] > inline_style_ranges[j + 1]['offset']:
                        inline_style_ranges[j], inline_style_ranges[j + 1] = inline_style_ranges[j + 1], \
                                                                             inline_style_ranges[j]

            saves = {}  # in there saves the run, to if have a new style for old word, the program can set this
            # Will split block text to letter, and set styles for all letters
            for word_i in range(len(block['text'])):
                run = paragraph.add_run(block['text'][word_i])  # create run for current letter
                run.font.size = self.__font_size
                run.font.name = self.__font_name
                run.font.color.rgb = self.__font_color_rgb
                saves.update(
                    {word_i: {'run': run, 'text': block['text'][word_i]}})  # and save that run and text to saves
            for styling in inline_style_ranges:  # set styles
                for word_i in range(styling['offset'],
                                    styling['offset'] + styling['length']):  # create 'for' in the letters (and styles)
                    if word_i in saves:
                        run = saves[word_i]['run']  # get current run for letter
                        run = set_styles(styling['style'], run, self=self,
                                         text=saves[word_i]['text'])  # get style of letter
                        saves.update({word_i: {'run': run, 'text': block['text'][
                            word_i]}})  # save new changes in the saves and in the run
        except Exception:
            raven.captureException()

    def _add_link_by_run(self, run_old, href, text=''):
        '''
        :param run_old: run objects where need to append link
        :param href: href
        :param text: text of link
        :return: new run, where have link
        '''
        try:
            r_id = self._paragraph.part.relate_to(href, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = oxml.shared.OxmlElement('w:hyperlink')  # create xml objects of hyperlkink
            hyperlink.set(oxml.shared.qn('r:id'), r_id, )  # add r_id

            # Create a w:r element and a new w:rPr element
            new_run = oxml.shared.OxmlElement('w:r')  # create run objects
            rPr = oxml.shared.OxmlElement('w:rPr')  # return new run of paragraph

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            run_old.text = ''
            run = run_old
            run._r.append(hyperlink)
            return run
        except Exception:
            raven.captureException()

    def _add_link(self, text: str, link: str, header=False):

        try:

            r_id = self._paragraph.part.relate_to(link, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = oxml.shared.OxmlElement('w:hyperlink')  # create xml objects of hyperlkink
            hyperlink.set(oxml.shared.qn('r:id'), r_id, )  # add r_id

            # Create a w:r element and a new w:rPr element
            new_run = oxml.shared.OxmlElement('w:r')  # create run objects
            rPr = oxml.shared.OxmlElement('w:rPr')  # return new run of paragraph

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            run = self._paragraph.add_run()
            run._r.append(hyperlink)
            # header uses when need to create link of image
            if header:
                self._paragraph.alignment = 1
                run.font.name = 'System'
                run.font.size = Pt(14)
            return run
        except Exception:
            raven.captureException()

    def _add_list_item(self, block):
        try:
            depth_level = block.get('depth', 1)
            if block['type'] == 'ordered-list-item':  # if need o add ordered list item
                depth_level = depth_level + 2 if depth_level + 2 <= 4 else 3  # correcting depth. It can't be much then 3
                style_now = 'List Number {depth_level}'.format(depth_level=depth_level)

                if self._paragraph.style.name != style_now:  #
                    self.__add_paragraph(
                        self.document)  # That is needs for word, to it shows that two lists like differents
                    self._paragraph.add_run('\n')  #
                self.__add_paragraph(self.document, style_now)  # added depth in to the style
            elif block['type'] == 'unordered-list-item':  # if list item is unordered
                depth_level += 2
                self.__add_paragraph(self.document, style='List Bullet {depth_level}'.format(
                    depth_level=depth_level))  # added depth in to the style
            self.add_text(block)  # add text of lists
        except Exception:
            raven.captureException()

    def _add_blockquote(self, block, style, position=None, reset_non_general=True):
        try:
            self.__depth_level += block['depth']

            if reset_non_general:
                self.__reset_non_general()  # reset values

            if not position:
                position = self.document  # get position (document)

            if style == 'warning':  # if style of blockquote item is warning
                color = RGBColor(192, 80, 77)  # set red color
                shadow_color = 'FFD9D9'
                blockquote_type = '🚨'  # set icon warning
                style_type = 'warning'  # it needs for creating style in docx
                border_color = 'FF0000'
                img = 'docsie/converters/converters/source/warning.png'
            elif style == 'question':  # if style of blockquote item is question
                color = RGBColor(128, 100, 162)  # set purple color
                blockquote_type = '❓'  # set question icon
                style_type = 'question'  # it needs for creating style in docx
                shadow_color = 'E7FDF8'
                border_color = '00CC00'
                img = 'docsie/converters/converters/source/question.png'
            else:  # style is default
                color = RGBColor(79, 129, 189)  # set light blue color
                blockquote_type = 'ℹ️'
                style_type = 'default'  # it needs for creating style in docx
                shadow_color = 'EAEEF2'
                border_color = '404040'
                img = 'docsie/converters/converters/source/info.png'
            try:
                self.document.styles.get_style_id(f'{style_type}docsie', WD_STYLE_TYPE.PARAGRAPH)
                # trying to get a style f'{style_type}docsie'.
                # if we will get a KeyError exception, this means that style of this blockquote isn't added.
                # we need to add it
            except KeyError:
                self.document.styles.add_style(f'{style_type}docsie', WD_STYLE_TYPE.PARAGRAPH)

            self.__add_paragraph(position, style='Normal')
            # style needs be added by self.document.styles.add_style



            table = self.document.add_table(rows=1, cols=2)
            table.autofit = False

            text_cell = table.rows[0].cells[1]
            text_cell.width = Inches(6.0)

            quote_cell = table.rows[0].cells[0]
            quote_cell.width = Inches(0.50)

            self.add_text(block=block, paragraph=text_cell.paragraphs[0])

            paragraph = table.rows[0].cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img, width=Inches(0.35), height=Inches(0.35))
            paragraph.style = self.document.styles[f'{style_type}docsie']
            quote_cell.paragraphs[0].style = self.document.styles[f'{style_type}docsie']
            text_cell.paragraphs[0].style = self.document.styles[f'{style_type}docsie']
            try:
                text_cell.paragraphs[0].runs[0].font.size = self.__font_size
            except IndexError:
                pass

            self.set_cell_border(quote_cell,
                                 start={'sz': 12, 'color': '#' + border_color, 'val': 'single'}, )
            text_shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shadow_color))
            quote_shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shadow_color))
            text_cell._tc.get_or_add_tcPr().append(text_shading)
            quote_cell._tc.get_or_add_tcPr().append(quote_shading)

            self.__depth_level -= block['depth']  # remove depth
        except Exception:
            raven.captureException()

    def _add_header_step(self, block, position=None, ):  # function to add header step
        try:
            self.__depth_level = block['depth']  # set depth level
            if not position:
                position = self.document  # set position (document)
            self.__add_paragraph(position)  # add paragraph
            self.step_header_elems['num'] += 1  # update header step number
            num = str(self.step_header_elems['num']) + '. '  # add prefix of header step
            self.add_text({'text': num, 'type': 'unstyled', 'offset': 0, 'length': len(num), 'entityRanges': [],
                           'inlineStyleRanges': [
                               {'offset': 0, 'length': len(num), 'style': 'header-step'}]})  # add header step prefix
            self.add_text(block)  # add text of header step
        except Exception:
            raven.captureException()


    def _add_table_cell(self, block):  # function to add cell\
        try:
            if self.doc_version == 3:
                if block.get('data', {}).get('table'):
                    self.__table = self.document.add_table(rows=1, cols=block['data']['table']['cols'])
                    self.__table.style = 'Table Grid'
                    self.__table.autofit = False

                    self.cells = 0  # reset cells
            else:
                index_old = self.keys.index(block['key']) - 1  # get old index of block
                old_block = self.article_blocks[index_old]
                if old_block['type'] != 'cell' or index_old < 0 or old_block['depth'] != block['depth']:
                    self.__table = None
                if not self.__table:  # if not have table, need to create a new table
                    if block['depth'] == 0:  # if depth is 0, that means that this table have 2 columns
                        self.__table = self.document.add_table(rows=1, cols=2)
                        self.__table.style = 'Table Grid'
                        self.__table.autofit = False
                    elif block['depth'] == 1:  # if depth is 1, that means that this table have 4 columns
                        self.__table = self.document.add_table(rows=1, cols=4)
                        self.__table.style = 'Table Grid'
                        self.__table.autofit = False
                    elif block['depth'] == 3:  # if depth is 3, that means that this table have 3 columns
                        self.__table = self.document.add_table(rows=1, cols=3)
                        self.__table.style = 'Table Grid'
                        self.__table.autofit = False
                    elif block['depth'] == 2:  # if depth is 2, that means that this table have 2 columns
                        self.__table = self.document.add_table(rows=1, cols=2)
                        self.__table.style = 'Table Grid'
                        self.__table.autofit = False
                    else:  # set to default 1 column
                        self.__table = self.document.add_table(rows=1, cols=1)
                        self.__table.style = 'Table Grid'
                        self.__table.autofit = False

            last_row = len(self.__table.rows) - 1  # this is last row of table
            row = self.__table.rows[last_row]
            try:
                row.cells[self.cells]  # try to use in the row last cell
            except IndexError:  # that means that need to reset cells, and add a new row
                self.cells = 0  # reset cells
                self.__table.add_row()  # add a new row
                row = self.__table.rows[len(self.__table.rows) - 1]  # set new row in the 'row'
            self.add_text(block, paragraph=row.cells[self.cells].paragraphs[0])
            self.cells += 1  # update cells
        except Exception:
            raven.captureException()

    def _add_dict(self, block):  # function to add dictionary
        '''
        this function work like _add_table_cell, the only difference is that this function set once 2 columns.
        '''
        try:
            index_old = self.keys.index(block['key']) - 1  # get old index of block
            old_block = self.article_blocks[index_old]  # get old block
            if old_block[
                'type'] != 'dictionary' or 0 > index_old:  # if old block isn't dictionary,old index is goes beyond of blocks, create a new dictionary
                self.__dict = None  # reset dictionary
            if not self.__dict:  # create a new dictionary, if dictinary is reseted
                self.__dict = self.document.add_table(rows=1, cols=3)  # create a new table
                self.__dict.style = 'Table Grid'
                self.__dict.autofit = False
                self.__dict.allow_autofit = False
                self.__dict.rows.table.columns[1].width = Inches(0.17)
                self.cells_d = 0
                self.rows_d = 0
            last_row = len(self.__dict.rows) - 1  # get last row index from table
            row = self.__dict.rows[last_row]  # get last row

            try:
                row.cells[self.cells_d]  # try get cell
            except IndexError:  # if it returned IndexError, that means to need to create a new row
                self.cells_d = 0  # reset cells
                self.__dict.add_row()  # create a new row
                last_row = len(self.__dict.rows) - 1  # get last row index from table
                row = self.__dict.rows[last_row]  # set in row a new row

            if last_row % 2 == 0:
                for cell in row.cells:
                    shading_elm = parse_xml('<w:shd {} w:fill="E7E7F9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

            if self.cells_d == 1:
                row.cells[self.cells_d].width = Inches(0.3)
                self.cells_d = 2  # skipping adding text to dividing column

            self.add_text(block, paragraph=row.cells[self.cells_d].paragraphs[0])
            self.cells_d += 1  # update index of cell
        except Exception:
            raven.captureException()

    def __indent_table(self, indent):
        try:
            tbl_pr = self.__table._element.xpath('w:tblPr')
            if tbl_pr:
                e = OxmlElement('w:tblInd')
                e.set(qn('w:w'), str(indent))
                e.set(qn('w:type'), 'dxa')
                tbl_pr[0].append(e)
        except Exception:
            raven.captureException()

    def add_ToC_element(self):  # that function create a 'Table of contents' element
        try:
            self.__add_paragraph(self.document)
            run = self._paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')  # creates a new element
            fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
            instrText.text = r'TOC \o "1-3" \h \z \u'  # change 1-3 depending on heading levels you need

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:t')
            fldChar3.text = 'Right-click to update field.'
            fldChar2.append(fldChar3)

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(fldChar4)
            self._paragraph.add_run('\nUpdate the Table of Contents to see the new contents of document ⬆️')
        except Exception:
            raven.captureException()

    def _set_picture(self, url, label='', depth=0, reset_non_general=True, position=None, set_automaticly=True,
                     alignment=0, border: bool = False):
        try:
            if not url:
                return  # some times url is ''
            self.__depth_level += depth

            if reset_non_general:
                self.__reset_non_general()  # reset values

            if not position:
                position = self.document

            self.__add_paragraph(position)
            self._paragraph.alignment = alignment

            crap_list = ['%', '#', '+', '@']
            basename = '{}/{}/images/'.format(self.shelf['id'], self.shelf['request_user_id'])

            if 'data:image' in url:
                fp = BytesIO()
                header, content = url.split(',')
                content = base64.b64decode(content)
                image = PIL_Image.open(BytesIO(content))
                image.save(fp, format=image.format)
                image_stream = BytesIO(fp.getvalue())
                filename = basename + hashlib.md5(content).hexdigest() + '.' + header.split('/')[1].split(';')[0]
                label = ''
            else:
                filename = url.split('/')[-1]
                filename = remove_crap(crap_list, filename)
                filename = basename + filename
                response = requests.get(url)
                content = response.content
                if response.headers['Content-Type'] == 'image/svg+xml':
                    if 'svg2png' in globals():
                        content = svg2png(content)
                        border = False
                        label = ''
                        # users in svg saving logos or something like that.
                        # Logos looks better without labels and borders
                    else:
                        raise ValueError(
                            'Do not have installed https://doc.courtbouillon.org/weasyprint/stable/first_steps.html')

                if border:
                    rgb_color_of_border = (233, 240, 255)
                    image_stream_for_change = BytesIO(content)
                    image = PIL_Image.open(image_stream_for_change)

                    master_image_size = image.size

                    width_of_border = int(master_image_size[0] // 50)

                    image_format = image.format
                    if not image_format:
                        image_format = 'PNG'

                    border_image_size = (master_image_size[0] + width_of_border, master_image_size[1] + width_of_border)
                    border_image = PIL_Image.new('RGB', border_image_size, rgb_color_of_border)
                    border_image.paste(image, ((border_image_size[0] - master_image_size[0]) // 2,
                                               (border_image_size[1] - master_image_size[1]) // 2))

                    image = border_image

                    image_stream = BytesIO()

                    image.save(image_stream, format=image_format)
                else:
                    image_stream = BytesIO(content)
            if set_automaticly:
                try:
                    figure = self._paragraph.add_run().add_picture(image_stream)
                except UnrecognizedImageError as e:
                    raise e
                if figure.width > self.images_max_width:
                    aspect_ratio = float(figure.width) / float(figure.height)
                    figure.width = self.images_max_width
                    figure.height = int(self.images_max_width / aspect_ratio)
                if label:
                    self.__add_paragraph(self.document, style='Caption')
                    run = self._paragraph.add_run(label)
                    run.font.color.rgb = self.__font_color_rgb
                    set_bold(run)

            else:
                return image_stream

            self.__depth_level -= depth
        except Exception:
            raven.captureException()

    def __reset_non_general(self):
        try:
            self.__table = None
            self.__table_col_index = 0
            self.__cols_num = 0
        except Exception:
            raven.captureException()

    def _add_figure(self, block):
        try:
            position = block['data'].get('align', '')
            if position == 'center':
                align = 1
            elif position == 'right':
                align = 2
            elif position == 'left':
                align = 0
            else:
                align = 1
            if 'data' in block:
                src = block['data'].get('src', '')
                label = block['data'].get('label', '')
                depth = block['data'].get('depth', 0)
                self._set_picture(depth=depth,
                                  reset_non_general=True,
                                  url=src,
                                  label=label,
                                  alignment=align,
                                  border=True
                                  )
        except Exception:
            raven.captureException()

    def save(self, shelf_id, request_user_id, filename):  # upload .docx document to server
        try:
            content = BytesIO()  # create empty document
            self.document.save(content)  # save to the empty document

            bucket = settings.DOCSIE_UPLOAD_BUCKET_NAME
            location = 'export/{}/{}/{}'.format(shelf_id, request_user_id,
                                                filename)  # create location of document for server
            resource = settings.S3_RESOURCE
            s3_client = settings.S3_CLIENT
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # set content type of .docx

            obj = resource.Object(bucket, location)
            obj.put(Body=content.getvalue(), ContentType=content_type)  # upload to server
            # get url
            url = s3_client.generate_presigned_url('get_object',
                                                   Params={'Bucket': bucket,
                                                           'Key': location},
                                                   ExpiresIn=900)

            return url  # return url
        except Exception:
            raven.captureException()

    def save_to_file(self,
                     instance: Union[Book, Documentation],
                     user: User,
                     filename: str
                     ) -> File:
        content = BytesIO()  # create empty document
        self.document.save(content)  # save to the empty document

        kwarg_map = {
            Book: 'book',
            Documentation: 'documentation'
        }
        kwarg = kwarg_map[instance._meta.model]
        kwargs = {'user': user, kwarg: instance}
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        upload_provider = DocumentUploadProvider.from_content(filename,
                                                              content.getvalue(),
                                                              content_type,
                                                              **kwargs)
        return upload_provider.upload()


def upload_image_to_bucket(filename, image_content, content_type):
    bucket = settings.DOCSIE_UPLOAD_BUCKET_NAME
    # TODO make sure location has doc_id/user_id/ in path
    location = 'export/' + filename
    resource = settings.S3_RESOURCE
    s3_client = settings.S3_CLIENT
    obj = resource.Object(bucket, location)
    response = obj.put(Body=image_content, ACL='public-read', ContentType=content_type)

    url = os.path.join(settings.KEYCDNURL, location)

    # TODO remove image from bucket

    return 'http://' + url


def delete_image_from_bucket(url):
    bucket = settings.DOCSIE_UPLOAD_BUCKET_NAME
    location = url.split(settings.KEYCDNURL)[1]
    location = location[1:] if location[0] == '/' else location
    resource = settings.S3_RESOURCE

    resource.Object(bucket, location).delete()


class DocxMergeStyles():
    '''
    That class will merge styles from second document into first document.
    First document - the input document, document in which need merge styles (input_doc)
    Second document - the document which will used for merge styles into first document (based_doc)
    Text and content will save in input_doc, that class only change styles and font.
    '''

    def __init__(self, copy_attrs: list = []):
        '''
        :param copy_attrs: attributes which need to merge
        '''
        try:
            if not copy_attrs:
                self.__copy_attrs = [
                    'italic',
                    'math',
                    'no_proof',
                    'web_hidden',
                    'strike',
                    'subscript',
                    'rtl',
                    'size',
                    'color.rgb',
                    'shadow',
                    'highlight_color',
                    'hidden',
                    'cs_bold',
                    'name',
                ]
            else:
                self.__copy_attrs = copy_attrs

        except Exception:
            raven.captureException()

    def __merge(self, input_doc: Document, based_doc: Document) -> Document:
        '''
        Basically function for merge styles of documents.
        :param input_doc: Document in which need merge styles
        :param based_doc: Document from which need copy styles and merge into input_doc
        :return: input_doc which has styles from based_doc
        '''
        try:
            based_fonts = {}  # initialize a variable which will save fonts from based_doc

            # 'for' which save in based_fonts basically fonts
            for pr in based_doc.paragraphs:

                for run in pr.runs:
                    if getattr(run, 'font', None):
                        based_fonts.update({pr.style.name: run.font})
                    # break

            for paragraph in input_doc.paragraphs:
                if based_fonts.get(paragraph.style.name):
                    based_style = based_fonts.get(paragraph.style.name)
                    for attr in self.__copy_attrs:  # if font is deep (for example font.color.rgb)
                        if '.' in attr:

                            pr = reduce(getattr, attr.split('.')[:-1],
                                        paragraph.style.font)  # get paragraph pre last font attribute

                            bs = reduce(getattr, attr.split('.'),
                                        based_style)  # get attribute from based document fonts

                            setattr(pr, attr.split('.')[-1:][0],
                                    bs)  # set font attribute from based document into input document

                        else:  # if font is normal and doesn't deep we can apply that only by setattr
                            setattr(paragraph.style.font, attr, getattr(based_style, attr))
            return input_doc  # return the result of merge
        except Exception:
            raven.captureException()

    @staticmethod
    def merge(input_doc: Document, based_doc: Document):
        '''
        This is function for default merge.
        You can use only DocxMergeStyles.merge(input_doc, based_doc) and this will return to you the result of merge
        :param input_doc: Document in which need merge styles
        :param based_doc: Document from which need copy styles and merge into input_doc
        :return: the result of merge (input_doc which has styles from based_doc)
        '''
        try:
            return DocxMergeStyles().__merge(input_doc=input_doc, based_doc=based_doc)
        except Exception:
            raven.captureException()
